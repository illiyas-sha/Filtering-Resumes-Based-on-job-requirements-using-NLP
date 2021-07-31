import docx     #a Python library for creating and updating Microsoft Word (.docx) files
import re       #regular expression


def get_content_as_string(document_path):

    req_doc = docx.Document(document_path)
    fullText = ''


    #String formed by concatenating the text of each run in the paragraph. Tabs and line breaks in the XML are mapped to \t and \n characters respectively.
    for para in req_doc.paragraphs:
        fullText = fullText + ' ' + para.text

    tbl = req_doc.tables
    for table in req_doc.tables:
        for row in table.rows:
            rowText = ''
            for cell in row.cells:
                rowText = rowText + ' ' + cell.text
            fullText = fullText + ' '+ rowText
    fullText = re.sub('\W+', ' ', fullText) #Select only alpha numerics
    fullText = re.sub('[^A-Za-z]+', ' ', fullText) #select only alphabet characters
    fullText = fullText.lower()

    return fullText

